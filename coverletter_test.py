import re
import os
import sys
# import argparse
from docx import Document
from datetime import date

try:
    arg = str(sys.argv)
    company = sys.argv[1]
    state = sys.argv[2]
except IndexError:
    raise SystemExit(f"Usage: {str(sys.argv)} \nExpected: command line args: company state")
print(arg[::])

# print ('Number of arguments:', len(sys.argv), 'arguments.')
# print ('Argument List:', str(sys.argv))

# print (sys.argv[1], sys.argv[2])
print(company, state)

today = date.today()
# Textual month, day and year
date = today.strftime("%B %d, %Y")
print(date, "\n**\n")

def replace_string(filename):
    doc = Document(filename)
    for p in doc.paragraphs:
        if '${company}' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${company}' in inline[i].text:
                    text = inline[i].text.replace('${company}', company)
                    inline[i].text = text
            print (p.text)
        
        if '${date}' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${date}' in inline[i].text:
                    text = inline[i].text.replace('${date}', date)
                    inline[i].text = text
            print (p.text)

        if '${state}' in p.text:
            # print("FOUND STATE")
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${state}' in inline[i].text:
                    text = inline[i].text.replace('${state}', state)
                    inline[i].text = text
            print (p.text)

    # for i, para in enumerate(doc.paragraphs):
    #     print (i, para.text)

replace_string("/Users/dhruv/Documents/NYU CoverLetter/coverletterdraft.docx")
