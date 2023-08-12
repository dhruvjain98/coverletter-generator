import re
import os
import sys
from docx import Document

try:
    # arg = sys.argv[1]
    arg = str(sys.argv)
except IndexError:
    raise SystemExit(f"Usage: {sys.argv[0]} <>")
print(arg[::])

print ('Number of arguments:', len(sys.argv), 'arguments.')
print ('Argument List:', str(sys.argv))

# print ("Hello")
document = Document("/Users/dhruv/Documents/NYU CoverLetter/coverletterdraft.docx")
# paragraph = document.paragraphs[0]
# print(paragraph.text)

# 3 ${date} ${company} ${state} 
# 5 If you are looking for Software Engineers... ${company}.  
# 14 Aside from this, I have good communication... ${company}

# def docx_replace_regex(doc_obj, regex , replace):

#     for p in doc_obj.paragraphs:
#         if regex.search(p.text):
#             inline = p.runs
#             # Loop added to work with runs (strings with same style)
#             for i in range(len(inline)):
#                 if regex.search(inline[i].text):
#                     text = regex.sub(replace, inline[i].text)
#                     inline[i].text = text

# for i, para in enumerate(document.paragraphs):
#     print (i, para.text)

def replace_string(filename):
    doc = Document(filename)
    for p in doc.paragraphs:
        if '${company}' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${company}' in inline[i].text:
                    text = inline[i].text.replace('${company}', 'ZS')
                    inline[i].text = text
            print (p.text)
        
        if '${date}' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${date}' in inline[i].text:
                    text = inline[i].text.replace('${date}', 'Today')
                    inline[i].text = text
            print (p.text)

        if '${state}' in p.text:
            # print("FOUND STATE")
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '${state}' in inline[i].text:
                    text = inline[i].text.replace('${state}', 'NY')
                    inline[i].text = text
            print (p.text)

# for i, para in enumerate(document.paragraphs):
    # print (i, para.text)

replace_string("/Users/dhruv/Documents/NYU CoverLetter/coverletterdraft.docx")

# if __name__ == "__main__":
#     try:
#         print(len(argv))
#         arg1 = sys.argv[1]
#         company = sys.argv[2]
#         state = sys.argv[3]
#     except IndexError:
#         print ("Usage: " + os.path.basename(__file__) + " <arg1>")
#         print ("Usage: " + os.path.basename(__file__) + " <company>")
#         print ("Usage: " + os.path.basename(__file__) + " <state>")
#         sys.exit(1)
#     except ValueError:     
#         raise MyError, "Not enough arguments"

#     # start the program
#     program(arg1)